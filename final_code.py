import os
import re
import pdfplumber
import pandas as pd
from docx import Document
from sentence_transformers import SentenceTransformer, util

from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_community.vectorstores import Chroma


# ========= CONFIG =========
ESCO_PATH = "skills_en.csv"          # ESCO CSV in same folder
OPENAI_MODEL = "gpt-4o-mini"         # or any other OpenAI chat model
OUTPUT_DOCX = "final_resume.docx"    # output Word file name


# ========= BASIC HELPERS =========
def clean_lines(text: str):
    if not text:
        return []
    return [l.strip() for l in re.split(r"[\r\n]+", str(text)) if len(l.strip()) > 2]


def read_file(path: str) -> str:
    path_lower = path.lower()
    if path_lower.endswith(".pdf"):
        with pdfplumber.open(path) as pdf:
            pages = [p.extract_text() or "" for p in pdf.pages]
        return "\n".join(pages)
    elif path_lower.endswith(".docx"):
        doc = Document(path)
        return "\n".join(p.text for p in doc.paragraphs)
    else:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()


# ========= ESCO + MODEL LOADING =========
def load_esco_and_model():
    if not os.path.exists(ESCO_PATH):
        raise FileNotFoundError(
            f"{ESCO_PATH} not found. Put skills_en.csv in this folder."
        )

    df = pd.read_csv(ESCO_PATH)
    if "preferredLabel" not in df.columns or "conceptUri" not in df.columns:
        raise ValueError("skills_en.csv must contain 'preferredLabel' and 'conceptUri' columns.")

    labels = df["preferredLabel"].dropna().tolist()
    uris = dict(zip(df["preferredLabel"], df["conceptUri"]))

    print("⏳ Loading sentence-transformers model...")
    model = SentenceTransformer("all-MiniLM-L6-v2")

    print("⏳ Encoding ESCO skills (one-time per run)...")
    esco_emb = model.encode(labels, convert_to_tensor=True)

    return model, labels, esco_emb, uris


# ========= SKILL NORMALIZATION =========
def normalize_skills(text, model, labels, esco_emb, uris, threshold: float = 0.45):
    phrases = clean_lines(text)
    if not phrases:
        return []

    cand_emb = model.encode(phrases, convert_to_tensor=True)
    hits = util.semantic_search(cand_emb, esco_emb, top_k=1)

    normalized = []
    for i, hit_list in enumerate(hits):
        best = hit_list[0]
        if best["score"] >= threshold:
            skill = labels[best["corpus_id"]]
            normalized.append(
                {
                    "raw": phrases[i],
                    "esco": skill,
                    "uri": uris.get(skill, ""),
                    "score": round(float(best["score"]), 2),
                }
            )

    return normalized


# ========= ATS SCORE + MISSING SKILLS =========
def compute_ats_score(resume_skills, jd_skills):
    r = {s["esco"] for s in resume_skills}
    j = {s["esco"] for s in jd_skills}

    overlap = r & j
    missing = list(j - r)

    if not j:
        return 0.0, missing

    skill_score = len(overlap) / len(j)  # coverage

    if overlap:
        overlap_scores = [
            s["score"] for s in resume_skills if s["esco"] in overlap
        ]
        avg_confidence = sum(overlap_scores) / len(overlap_scores)
    else:
        avg_confidence = 0.0

    final_score = round((0.7 * skill_score + 0.3 * avg_confidence) * 100, 2)
    return final_score, missing


# ========= RAG-HAT REWRITER =========
def run_rag_hat(resume_text: str, jd_text: str, llm: ChatOpenAI, embeddings: OpenAIEmbeddings, top_k: int = 3):
    facts = clean_lines(resume_text)
    requirements = clean_lines(jd_text)

    if not facts or not requirements:
        print("⚠️ Not enough content for RAG-HAT (facts or JD empty).")
        return []

    vectordb = Chroma.from_texts(facts, embedding=embeddings)
    retriever = vectordb.as_retriever(search_kwargs={"k": top_k})

    optimized_bullets = []

    for req in requirements:
        docs = retriever.invoke(req)
        context = "\n".join([d.page_content for d in docs])
        prompt = f"""
You are a strict ATS resume rewriter.

RULES:
1. You may ONLY use the Resume Evidence below.
2. You are FORBIDDEN from inventing tools, companies, or metrics.
3. If evidence is weak, rewrite conservatively.
4. Output only bullet points.

Job Requirement:
{req}

Resume Evidence:
{context}

Generate 2 ATS-optimized bullets:
"""
        resp = llm.invoke(prompt)
        text = resp.content if hasattr(resp, "content") else str(resp)
        optimized_bullets.append(text.strip())

    # clean up vectorstore
    vectordb.delete_collection()
    return optimized_bullets


# ========= WORD EXPORT =========
def export_to_word(
    resume_text: str,
    jd_text: str,
    ats_score: float,
    missing_skills,
    optimized_bullets,
    output_path: str = OUTPUT_DOCX,
):
    doc = Document()

    # Header
    doc.add_heading("ATS Optimized Resume", level=1)
    doc.add_paragraph(f"Final ATS Score: {ats_score}/100")

    # Missing skills
    doc.add_heading("Missing Skills (From JD)", level=2)
    if missing_skills:
        for m in missing_skills:
            doc.add_paragraph(m, style="List Bullet")
    else:
        doc.add_paragraph("No critical missing skills identified based on ESCO matching.")

    # RAG-HAT bullets
    doc.add_heading("RAG-HAT Optimized Experience Bullets", level=2)
    if optimized_bullets:
        for block in optimized_bullets:
            for line in block.split("\n"):
                if line.strip():
                    doc.add_paragraph(line.strip(), style="List Bullet")
    else:
        doc.add_paragraph("No optimized bullets generated (insufficient input).")

    # Optional: add JD raw text for reference
    doc.add_heading("Job Description (Reference)", level=2)
    for line in clean_lines(jd_text):
        doc.add_paragraph(line)

    # Optional: add original resume raw extract at end
    doc.add_heading("Original Resume (Extract)", level=2)
    for line in clean_lines(resume_text):
        doc.add_paragraph(line)

    doc.save(output_path)
    print(f"✅ Final optimized resume saved → {output_path}")


# ========= MAIN ORCHESTRATOR =========
def run_pipeline(resume_path: str, jd_path: str, output_docx: str = OUTPUT_DOCX):
    # Basic sanity checks
    assert os.path.exists(ESCO_PATH), f"❌ ESCO file not found: {ESCO_PATH}"
    assert os.path.exists(resume_path), f"❌ Resume file not found: {resume_path}"
    assert os.path.exists(jd_path), f"❌ JD file not found: {jd_path}"
    assert os.getenv("OPENAI_API_KEY"), "❌ OPENAI_API_KEY is not set in environment."

    print("📄 Reading resume...")
    resume_text = read_file(resume_path)

    print("📄 Reading job description...")
    jd_text = read_file(jd_path)

    print("📚 Loading ESCO + model...")
    model, labels, esco_emb, uris = load_esco_and_model()

    print("🔍 Normalizing resume skills...")
    resume_skills = normalize_skills(resume_text, model, labels, esco_emb, uris)

    print("🔍 Normalizing JD skills...")
    jd_skills = normalize_skills(jd_text, model, labels, esco_emb, uris)

    print("📊 Computing ATS score + missing skills...")
    ats, missing = compute_ats_score(resume_skills, jd_skills)
    print(f"   ▶ ATS Score: {ats}/100")
    print(f"   ▶ Missing skills: {missing}")

    print("🧠 Initializing LLM + embeddings for RAG-HAT...")
    llm = ChatOpenAI(model=OPENAI_MODEL, temperature=0.2)
    emb = OpenAIEmbeddings()

    print("🧠 Running RAG-HAT to optimize bullets...")
    optimized = run_rag_hat(resume_text, jd_text, llm, emb, top_k=3)

    print("📝 Exporting final Word document...")
    export_to_word(
        resume_text=resume_text,
        jd_text=jd_text,
        ats_score=ats,
        missing_skills=missing,
        optimized_bullets=optimized,
        output_path=output_docx,
    )


if __name__ == "__main__":
    # Change these to your actual filenames
    RESUME_PATH = r"C:\Users\gayat\OneDrive\Desktop\humanAI\ADV NLP\resmatch (1)\resume.pdf"   # or "resume.docx"
    JD_PATH = r"C:\Users\gayat\OneDrive\Desktop\humanAI\ADV NLP\resmatch (1)\jd.txt"

    run_pipeline(RESUME_PATH, JD_PATH, OUTPUT_DOCX)
